import streamlit as st
from docx import Document
from docx.shared import Inches
import io
import time

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        color: #4CAF50;
        text-align: center;
        font-weight: bold;
        margin-bottom: 2rem;
    }
    .sub-header {
        font-size: 1.5rem;
        color: #2196F3;
        margin-bottom: 1rem;
    }
    .input-section {
        background-color: #f9f9f9;
        padding: 1rem;
        border-radius: 10px;
        margin-bottom: 1rem;
    }
    .generate-btn {
        background-color: #4CAF50;
        color: white;
        font-size: 1.2rem;
        padding: 0.5rem 1rem;
        border: none;
        border-radius: 5px;
        cursor: pointer;
    }
</style>
""", unsafe_allow_html=True)

def generate_modul_ajar(data):
    doc = Document()

    # Header
    doc.add_heading('MODUL AJAR', 0)

    # I. IDENTITAS
    doc.add_heading('I. IDENTITAS', level=1)
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'

    fields = [
        ('Nama Sekolah', data['nama_sekolah']),
        ('Mata Pelajaran', data['mata_pelajaran']),
        ('Kelas/Semester', f"{data['kelas']}/{data['semester']}"),
        ('Materi Pokok', data['materi']),
        ('Alokasi Waktu', f"{data['alokasi_waktu']} JP"),
        ('Kompetensi Dasar', data['kompetensi_dasar']),
        ('Tujuan Pembelajaran', data['tujuan_pembelajaran']),
        ('Materi Pembelajaran', data['materi_pembelajaran'])
    ]

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value

    # II. KEGIATAN PEMBELAJARAN (Deskripsi Gamblang)
    doc.add_heading('II. KEGIATAN PEMBELAJARAN', level=1)

    # 1. Pendahuluan
    doc.add_heading('1. Kegiatan Pendahuluan (10 Menit)', level=2)
    p1 = doc.add_paragraph()
    p1.add_run("• Orientasi: ").bold = True
    p1.add_run("Guru membuka pelajaran dengan salam, doa bersama, dan memeriksa kehadiran siswa untuk mengondisikan suasana belajar yang positif.\n")
    p1.add_run("• Apersepsi: ").bold = True
    p1.add_run(f"Guru mengaitkan materi {data['materi']} dengan pengalaman siswa sehari-hari atau materi sebelumnya.\n")
    p1.add_run("• Motivasi: ").bold = True
    p1.add_run(f"Guru menyampaikan tujuan pembelajaran yang ingin dicapai dan manfaat mempelajari {data['materi']} dalam kehidupan.")

    # 2. Kegiatan Inti
    doc.add_heading('2. Kegiatan Inti (50 Menit)', level=2)
    p2 = doc.add_paragraph()
    p2.add_run("• Tahap 1: Pemberian Rangsangan (Stimulation)\n").bold = True
    p2.add_run(f"Siswa mengamati tayangan visual atau benda konkret yang berkaitan dengan konsep {data['materi']}.\n")
    
    p2.add_run("• Tahap 2: Identifikasi Masalah (Problem Statement)\n").bold = True
    p2.add_run("Guru memberikan pertanyaan pemantik untuk memicu rasa ingin tahu siswa dan diskusi awal secara klasikal.\n")
    
    p2.add_run("• Tahap 3: Pengumpulan Data & Kolaborasi (Data Collection)\n").bold = True
    p2.add_run(f"Siswa dibagi ke dalam kelompok heterogen untuk mengeksplorasi materi {data['materi']} melalui LKPD atau aktivitas praktik mandiri.\n")
    
    p2.add_run("• Tahap 4: Pengolahan Data & Komunikasi (Data Processing)\n").bold = True
    p2.add_run("Setiap kelompok mendiskusikan hasil temuan mereka dan menyajikannya dalam bentuk karya (tulisan/gambar/presentasi).\n")
    
    p2.add_run("• Tahap 5: Pembuktian & Refleksi (Verification)\n").bold = True
    p2.add_run("Guru memberikan penguatan (feedback) terhadap hasil presentasi kelompok untuk meluruskan miskonsepsi.")

    # 3. Penutup
    doc.add_heading('3. Kegiatan Penutup (10 Menit)', level=2)
    p3 = doc.add_paragraph()
    p3.add_run("• Menyimpulkan: ").bold = True
    p3.add_run(f"Siswa bersama guru menyimpulkan poin-poin utama dari pembelajaran materi {data['materi']}.\n")
    p3.add_run("• Evaluasi: ").bold = True
    p3.add_run("Guru melakukan penilaian singkat (post-test) atau refleksi perasaan siswa setelah belajar hari ini.\n")
    p3.add_run("• Tindak Lanjut: ").bold = True
    p3.add_run("Guru memberikan tugas pembiasaan atau menginformasikan materi untuk pertemuan berikutnya, lalu menutup dengan doa.")

    # III. PENILAIAN
    doc.add_heading('III. PENILAIAN', level=1)
    doc.add_paragraph(data['penilaian'])

    return doc

# Sidebar
st.sidebar.title("📚 Generator Modul Ajar")
st.sidebar.markdown("Aplikasi untuk membuat modul ajar secara otomatis.")
st.sidebar.markdown("---")
st.sidebar.markdown("**Fitur:**")
st.sidebar.markdown("- Input data mudah")
st.sidebar.markdown("- Generate dokumen Word")
st.sidebar.markdown("- Download instan")

# Main content
st.markdown('<div class="main-header">🎓 Generator Modul Ajar</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Buat modul ajar profesional dengan mudah!</div>', unsafe_allow_html=True)

# Input sections with expanders
with st.expander("📝 Identitas Modul", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        nama_sekolah = st.text_input("🏫 Nama Sekolah", placeholder="Masukkan nama sekolah")
        mata_pelajaran = st.text_input("📖 Mata Pelajaran", placeholder="Contoh: Matematika")
        kelas = st.text_input("👥 Kelas", placeholder="Contoh: X")
    with col2:
        semester = st.text_input("📅 Semester", placeholder="Contoh: 1")
        materi = st.text_input("📋 Materi Pokok", placeholder="Contoh: Aljabar")
        alokasi_waktu = st.number_input("⏰ Alokasi Waktu (JP)", min_value=1, value=2)

with st.expander("🎯 Kompetensi dan Tujuan"):
    kompetensi_dasar = st.text_area("🎯 Kompetensi Dasar", placeholder="Deskripsikan kompetensi dasar", height=100)
    tujuan_pembelajaran = st.text_area("🎯 Tujuan Pembelajaran", placeholder="Deskripsikan tujuan pembelajaran", height=100)

with st.expander("📚 Materi Pembelajaran"):
    materi_pembelajaran = st.text_area("📚 Materi Pembelajaran", placeholder="Deskripsikan materi pembelajaran", height=150)

with st.expander("📊 Penilaian"):
    penilaian = st.text_area("📊 Penilaian", placeholder="Deskripsikan metode penilaian", height=100)

# Generate button
if st.button("🚀 Generate Modul Ajar", key="generate"):
    if nama_sekolah and mata_pelajaran and kelas and semester and materi and kompetensi_dasar and tujuan_pembelajaran and materi_pembelajaran and penilaian:
        with st.spinner("Membuat modul ajar..."):
            time.sleep(1)  # Simulate processing
            data = {
                'nama_sekolah': nama_sekolah,
                'mata_pelajaran': mata_pelajaran,
                'kelas': kelas,
                'semester': semester,
                'materi': materi,
                'alokasi_waktu': alokasi_waktu,
                'kompetensi_dasar': kompetensi_dasar,
                'tujuan_pembelajaran': tujuan_pembelajaran,
                'materi_pembelajaran': materi_pembelajaran,
                'penilaian': penilaian
            }
            
            doc = generate_modul_ajar(data)
            
            # Save to BytesIO
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            
            st.success("✅ Modul ajar berhasil dibuat!")
            st.download_button(
                label="📥 Download Modul Ajar (.docx)",
                data=bio,
                file_name=f"modul_ajar_{materi.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download"
            )
    else:
        st.error("❌ Harap isi semua field yang diperlukan.")

# Footer
st.markdown("---")
st.markdown("Dibuat dengan ❤️ menggunakan Streamlit")cd /workspaces/modulajar
git add .
git commit -m "Add generator.py with modul ajar generation logic"
